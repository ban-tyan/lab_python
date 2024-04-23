import docx

def generate_doc_report(size, fabric_cost, total_cost):
    # Создание нового документа Word
    document = docx.Document()

    # Добавление заголовка
    document.add_heading("Расчет стоимости одежды", 0)

    # Добавление таблицы
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Размер"
    table.cell(0, 1).text = str(size)
    table.cell(1, 0).text = "Стоимость ткани"
    table.cell(1, 1).text = str(fabric_cost)
    table.cell(2, 0).text = "Общая стоимость"
    table.cell(2, 1).text = str(total_cost)

    # Сохранение документа
    document.save("stoimost.docx")  # Исправлено название файла
