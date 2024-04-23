import tkinter as tk
from modules.calc_info import *
from docx import Document  # Импорт библиотеки Document

class ClothingCostCalculator:
    def __init__(self):  # Исправил опечатку в названии метода
        self.window = tk.Tk()
        self.window.title("Расчет стоимости одежды")

        self.clothing_types = [Jacket(), Trousers(), Suit()]

        self.clothing_type_var = tk.StringVar()
        self.clothing_type_var.set("Одежда")  # Default 

        self.size_label = tk.Label(self.window, text="Размер:")
        self.size_label.pack()

        self.size_entry = tk.Entry(self.window)
        self.size_entry.pack()

        self.fabric_price_label = tk.Label(self.window, text="Цена ткани за кв. м.:")
        self.fabric_price_label.pack()

        self.fabric_price_entry = tk.Entry(self.window)
        self.fabric_price_entry.pack()

        self.clothing_type_menu = tk.OptionMenu(self.window, self.clothing_type_var, *["Пиджак", "Брюки", "Тройка"])
        self.clothing_type_menu.pack()

        self.calculate_button = tk.Button(self.window, text="Рассчитать", command=self.calculate_cost)
        self.calculate_button.pack()

        self.save_button = tk.Button(self.window, text="Сохранить в DOCX", command=self.save_to_docx)
        self.save_button.pack()

        self.result_label = tk.Label(self.window, text="")
        self.result_label.pack()

        self.window.mainloop()

    def calculate_cost(self):
        selected_clothing_type = self.clothing_types[["Пиджак", "Брюки", "Тройка"].index(self.clothing_type_var.get())]
        size = float(self.size_entry.get())
        fabric_price_per_square_meter = float(self.fabric_price_entry.get())

        total_cost = Calculations.calculate_total_cost(selected_clothing_type, size, fabric_price_per_square_meter)

        self.result_label.configure(text=f"Общая стоимость: {total_cost}")

    def save_to_docx(self):
        selected_clothing_type = self.clothing_types[["Пиджак", "Брюки", "Тройка"].index(self.clothing_type_var.get())]
        size = float(self.size_entry.get())
        fabric_price_per_square_meter = float(self.fabric_price_entry.get())
        total_cost = Calculations.calculate_total_cost(selected_clothing_type, size, fabric_price_per_square_meter)

        document = Document()
        document.add_heading("Расчет стоимости одежды", 0)

        # Добавление таблицы
        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "Тип одежды"
        table.cell(0, 1).text = selected_clothing_type.name
        table.cell(1, 0).text = "Размер"
        table.cell(1, 1).text = str(size)
        table.cell(2, 0).text = "Стоимость ткани"
        table.cell(2, 1).text = str(fabric_price_per_square_meter)
        table.cell(3, 0).text = "Общая стоимость"
        table.cell(3, 1).text = str(total_cost)

        # Сохранение документа
        document.save("stoimost.docx") 

